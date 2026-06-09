import io
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document
from fastapi import HTTPException, UploadFile


ALLOWED_EXTENSIONS = {".pdf", ".docx"}
MAX_FILE_SIZE = 5 * 1024 * 1024  # 5 MB


def _extract_pdf_text(content: bytes) -> str:
    doc = fitz.open(stream=content, filetype="pdf")
    try:
        parts = []
        for page in doc:
            text = page.get_text()
            if text.strip():
                parts.append(text.strip())
        return "\n\n".join(parts)
    finally:
        doc.close()


def _extract_docx_text(content: bytes) -> str:
    doc = Document(io.BytesIO(content))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())
    return "\n\n".join(parts)


async def parse_resume(file: UploadFile) -> tuple[str, str]:
    filename = file.filename or "resume"
    ext = Path(filename).suffix.lower()

    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type '{ext}'. Allowed: PDF, DOCX.",
        )

    content = await file.read()
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(status_code=400, detail="File exceeds 5 MB limit.")

    if not content:
        raise HTTPException(status_code=400, detail="Uploaded file is empty.")

    try:
        if ext == ".pdf":
            text = _extract_pdf_text(content)
        else:
            text = _extract_docx_text(content)
    except Exception as exc:
        raise HTTPException(
            status_code=422,
            detail=f"Could not parse resume: {exc}",
        ) from exc

    text = text.strip()
    if len(text) < 20:
        raise HTTPException(
            status_code=422,
            detail="Could not extract enough text from the resume.",
        )

    return text, filename
